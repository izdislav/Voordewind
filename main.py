# This is a sample Python script.

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import RGBColor

pd.set_option('display.width', 400)
pd.set_option('display.max_columns', 10)


def get_right_answer(row, i):
    j = -1
    for cell in row.cells:
        # print(cell.text)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.font.color.rgb == RGBColor(0x00, 0x00, 0xff):
                    return np.array([i, j])
        j = j + 1
    return


def read_table_with_colors(document):
    # tables = document.tables
    # for table in tables:
    input_array = np.array(['Номер', 'Въпрос', 'ОТГОВОР А', 'ОТГОВОР Б', 'ОТГОВОР В', 'ОТГОВОР Г'])

    table = document.tables[0]
    data = [[cell.text for cell in row.cells] for row in table.rows]
    df_questions = pd.DataFrame(data)
    i = 0
    answers = np.array([0, 0])
    for row in table.rows:
        right_answer = get_right_answer(row, i)
        if isinstance(right_answer, np.ndarray):
            answers = np.vstack((answers, right_answer))
        i = i + 1
    df_answers = pd.DataFrame(answers)
    print(df_questions)
    print(df_answers)


def read_docx_table(document, table_num=1, nheader=1):
    table = document.tables[table_num - 1]
    data = [[cell.text for cell in row.cells] for row in table.rows]
    df = pd.DataFrame(data)
    if nheader == 1:
        df = df.rename(columns=df.iloc[0]).drop(df.index[0]).reset_index(drop=True)
    return df


# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    document = Document("40BT.docx")
    read_table_with_colors(document)
    # dataframe = read_docx_table(document, 1, 0)
    # print(dataframe)

# See PyCharm help at https://www.jetbrains.com/help/pycharm/
