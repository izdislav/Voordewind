from io import BytesIO
from xml.dom import minidom

import numpy as np
import pandas as pd
import random
from docx import Document
from docx.shared import RGBColor
from flask import Flask, render_template, request, url_for
import zipfile
import os
import shutil

app = Flask(__name__)


def get_right_answer(row, i):
    j = -1
    for cell in row.cells:
        # print(cell.text)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.font.color.rgb == RGBColor(0x00, 0x00, 0xff):
                    return np.array([i, j])
        j = j + 1
    return


def get_document_tables(document, selected_table_indexes):
    return_tables = []
    for selected_table in selected_table_indexes:
        return_tables.append(document.tables[selected_table])
    return return_tables


def get_answers(tables):
    # tables = document.tables
    # for table in tables:
    # input_array = np.array(['Номер', 'Въпрос', 'ОТГОВОР А', 'ОТГОВОР Б', 'ОТГОВОР В', 'ОТГОВОР Г'])
    i = 1
    answers = {}
    for table in tables:
        for row in table.rows:
            right_answer = get_right_answer(row, int(row.cells[0].text))
            if isinstance(right_answer, np.ndarray):
                answers[right_answer[0]] = right_answer[1]
            i = i + 1
    return answers


def get_table_data(document, tables):
    data = []
    for table in tables:
        for row in table.rows:
            if row.cells[0].text:
                inner_cell_xml = row.cells[1]._tc.xml
                image_stream = get_image_from_xml(document, inner_cell_xml)
                image_src = ''
                if image_stream:
                    img_name = row.cells[0].text + ".jpg"
                    with open("static/images/" + img_name, "wb") as f:
                        f.write(image_stream.getbuffer())
                    image_src = 'static/images/' + img_name
                data.insert(
                    int(row.cells[0].text),
                    [
                        int(row.cells[0].text),
                        row.cells[1].text,
                        row.cells[2].text,
                        row.cells[3].text,
                        row.cells[4].text,
                        row.cells[5].text,
                        image_src
                    ]
                )
    return data


def get_image_from_xml(source_document, xml_code):
    """
        Returns the rId for an image in the *xml_code*

        :xml_code
            xml_code: [string] xml code from which to extract the image from
        :return
            image_stream: [BytesIO stream] the image to find
            None if no image exists in the xml_file

    """
    # Parse the xml code for the blip
    xml_parser = minidom.parseString(xml_code)

    items = xml_parser.getElementsByTagName('a:blip')

    # Check if an image exists
    if items:
        # Extract the rId of the image
        r_id = items[0].attributes['r:embed'].value

        # Get the blob of the image
        source_document_part = source_document.part
        image_part = source_document_part.related_parts[r_id]
        image_bytes = image_part._blob

        # Write the image bytes to a file (or BytesIO stream) and feed it to document.add_picture(), maybe:
        image_stream = BytesIO(image_bytes)

        return image_stream
    else:
        return None


@app.route("/", methods=['POST', 'GET'])
def questions():
    selected_tables = [47, 48, 49]
    hard_questions = [48, 49, 70, 100, 125, 151, 170, 178, 193, 304, 440, 424, 443, 503, 559, 560, 574, 577, 578, 694, 744]
    if request.method == 'POST':
        selected_tables = request.form.getlist("table", int)
    document = Document("40BT1.docx")
    tables = get_document_tables(document, selected_tables)
    answers = get_answers(tables)
    questions = get_table_data(document, tables)
    sample_questions = random.sample(questions, 5)

    return render_template(
        'questions.html',
        result=sample_questions,
        answers=answers,
        tables_length=len(document.tables),
        selected_tables=selected_tables
    )


@app.route('/results', methods=['POST', 'GET'])
def results():
    return render_template('results.html')


if __name__ == "__main__":
    app.run()
